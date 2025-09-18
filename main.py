import streamlit as st
import pandas as pd
from deep_translator import GoogleTranslator
from gensim.summarization import summarize
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile

# ---------------- UI ------------------
st.set_page_config(page_title="Text Translator & Summarizer", page_icon="📘", layout="centered")

st.markdown("""
    <h1 style='color:red; text-align:center;'>Text Translator & Summarizer</h1>
    <h3 style='text-align:center;'>Presented by <span style='color:red;'>Vaibhavi Zunzunkar</span> – JD College</h3>
""", unsafe_allow_html=True)

st.write("### 📌 Paste text below or upload a file (txt, pdf, docx)")

# ---------------- File Upload ------------------
uploaded_file = st.file_uploader("Upload File", type=["txt", "pdf", "docx"])
text_input = st.text_area("Or enter text manually")

text = ""

if uploaded_file is not None:
    if uploaded_file.type == "text/plain":
        text = uploaded_file.read().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        import PyPDF2
        reader = PyPDF2.PdfReader(uploaded_file)
        text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        text = " ".join([p.text for p in doc.paragraphs])
else:
    text = text_input

# ---------------- Options ------------------
option = st.radio("Choose Action:", ["Summarize", "Translate"])

if text:
    if option == "Summarize":
        try:
            result = summarize(text, ratio=0.3)
            if not result:
                result = "⚠️ Text too short for summarization. Please enter longer text."
        except:
            result = "⚠️ Could not summarize. Try with more text."
    else:
        lang = st.selectbox("Select Language", ["hi", "mr", "fr", "de", "es", "ta"])
        try:
            result = GoogleTranslator(source="auto", target=lang).translate(text)
        except Exception as e:
            result = f"⚠️ Translation error: {e}"

    st.subheader("✅ Result:")
    st.write(result)

    # ---------------- Download Options ------------------
    def save_to_word(content):
        doc = Document()
        doc.add_paragraph(content)
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_file.name)
        return tmp_file.name

    def save_to_pdf(content):
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        c = canvas.Canvas(tmp_file.name, pagesize=letter)
        text_obj = c.beginText(40, 750)
        for line in content.split("\n"):
            text_obj.textLine(line)
        c.drawText(text_obj)
        c.save()
        return tmp_file.name

    col1, col2 = st.columns(2)
    with col1:
        if st.download_button("📄 Download Word", open(save_to_word(result), "rb"), file_name="output.docx"):
            pass
    with col2:
        if st.download_button("📕 Download PDF", open(save_to_pdf(result), "rb"), file_name="output.pdf"):
            pass
